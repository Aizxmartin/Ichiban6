from docx import Document
from docx.shared import Inches
from adjustments import calculate_adjustments
import io
import pandas as pd
import numpy as np

def fmt(value, decimals=0, prefix="$"):
    if pd.isna(value):
        return "N/A"
    return f"{prefix}{value:,.{decimals}f}" if prefix else f"{value:,.{decimals}f}"

def generate_report(df, subject_info, real_avm=None, redfin_estimate=None, zillow_estimate=None):
    from market_adjustment_schema import schema

    adjusted_rows = []
    for _, row in df.iterrows():
        total_adj, adjusted_price, ag_diff = calculate_adjustments(row, subject_info, schema)
        if total_adj is None:
            continue
        adjusted_rows.append({
            "Address": row.get("Street Address", "N/A"),
            "Close Price": row.get("Close Price", 0),
            "Concessions": row.get("Concessions", 0),
            "AG SF": row.get("Above Grade Finished Area", 0),
            "AG Diff": ag_diff,
            "Total Adj": total_adj,
            "Adjusted Price": adjusted_price,
            "Adj PPSF": adjusted_price / row.get("Above Grade Finished Area", 1)
        })

    doc = Document()
    doc.add_heading("Market Valuation Report – Adjusted Comparison with Breakdown", 0)

    doc.add_paragraph(f"Subject Property\nAddress: {subject_info.get('Address')}\n"
                      f"Above Grade SF: {subject_info.get('Above Grade Finished Area')}\n"
                      f"Bedrooms: {subject_info.get('Bedrooms')}\n"
                      f"Bathrooms: {subject_info.get('Bathrooms')}")

    if adjusted_rows:
        doc.add_paragraph("Comparable Properties")
        table = doc.add_table(rows=1, cols=8)
        hdr = table.rows[0].cells
        hdr[0].text = "Address"
        hdr[1].text = "Close Price"
        hdr[2].text = "Concessions"
        hdr[3].text = "AG SF"
        hdr[4].text = "AG Diff"
        hdr[5].text = "Total Adj"
        hdr[6].text = "Adjusted Price"
        hdr[7].text = "Adj PPSF"

        for comp in adjusted_rows:
            row_cells = table.add_row().cells
            row_cells[0].text = str(comp["Address"])
            row_cells[1].text = fmt(comp["Close Price"])
            row_cells[2].text = fmt(comp["Concessions"])
            row_cells[3].text = str(comp["AG SF"])
            row_cells[4].text = str(int(comp["AG Diff"]))
            row_cells[5].text = fmt(comp["Total Adj"])
            row_cells[6].text = fmt(comp["Adjusted Price"])
            row_cells[7].text = fmt(comp["Adj PPSF"], 2)

        avg_price = np.nanmean([c["Adjusted Price"] for c in adjusted_rows])
        avg_ppsf = np.nanmean([c["Adj PPSF"] for c in adjusted_rows])

        doc.add_paragraph(f"\nValuation Summary\n"
                          f"Average Adjusted Price: {fmt(avg_price)}\n"
                          f"Average Price Per SF: {fmt(avg_ppsf, 2)}")

        # Online estimate average
        online_vals = [v for v in [real_avm, redfin_estimate, zillow_estimate] if v is not None]
        online_avg = sum(online_vals) / len(online_vals) if online_vals else None

        if online_avg is not None and not np.isnan(avg_price):
            doc.add_paragraph(
                f"Online Estimate Average: {fmt(online_avg)}\n"
                f"Recommended Market Range: {fmt(online_avg)} – {fmt(avg_price)}"
            )
        else:
            doc.add_paragraph("Recommended Market Range: N/A")

    else:
        doc.add_paragraph("No valid comps available for valuation.")

    doc.add_paragraph("\nMethodology\n"
                      "Close Price is the original sale price. Adjustments are based on AG SF differences.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer